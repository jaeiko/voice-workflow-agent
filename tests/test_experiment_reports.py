"""Experiment report service tests use isolated temporary SQLite files only."""

from __future__ import annotations

import json
import asyncio
import csv
import io
import os
import tempfile
import unittest
import unittest.mock
from pathlib import Path

from voice_workflow_agent.experiment_reports import ExperimentReportStore


class ExperimentReportStoreTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.store = ExperimentReportStore(
            Path(self.temporary.name) / "reports.sqlite"
        )

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def open(self):
        return self.store.open_report(
            session_id="session-test-1",
            protocol_id="candidate-a-curated-development-v1",
            protocol_title="Candidate A development fixture",
            protocol_revision="revision-test-1",
            protocol_sha256="6" * 64,
            readiness_status="analysis_required",
            development_only=True,
        )

    def test_one_report_per_session_and_idempotent_events(self):
        first = self.open()
        second = self.open()
        self.assertEqual(first["report_id"], second["report_id"])
        one = self.store.append_event(
            first["report_id"], event_key="turn-1-step",
            event_type="step_completed", step_id="step-1", step_label="1",
            payload={"operation_id": "operation-1"},
        )
        duplicate = self.store.append_event(
            first["report_id"], event_key="turn-1-step",
            event_type="step_completed", step_id="step-1", step_label="1",
            payload={"operation_id": "operation-1"},
        )
        self.assertTrue(one["event_inserted"])
        self.assertFalse(duplicate["event_inserted"])
        self.assertEqual(len(duplicate["events"]), 1)

    def test_anomaly_and_finalization_are_counted_exactly_once(self):
        report = self.open()
        for _ in range(2):
            report = self.store.append_event(
                report["report_id"], event_key="anomaly-turn-2",
                event_type="anomaly", step_id="step-4", step_label="4",
                user_wording="예상과 다르게 색이 남아 있어.",
                category="protocol_block", severity="unknown",
                confirmation_state="user_reported",
            )
        self.assertEqual(report["anomaly_count"], 1)
        finalized = self.store.finalize(
            report["report_id"], status="stopped", event_key="stop-1"
        )
        repeated = self.store.finalize(
            report["report_id"], status="stopped", event_key="stop-1"
        )
        self.assertEqual(finalized["finalization_version"], 1)
        self.assertEqual(repeated["finalization_version"], 1)
        self.assertEqual(repeated["status"], "stopped")

    def test_markdown_and_json_exports_exclude_hidden_reasoning(self):
        report = self.open()
        self.store.append_event(
            report["report_id"], event_key="source-turn-3",
            event_type="source_consulted", source_tier="approved_lab_corpus",
            citation_identities=("a" * 64,), payload={"summary": "cited source"},
        )
        exported = json.loads(self.store.export_json(report["report_id"]))
        markdown = self.store.export_markdown(report["report_id"]).decode()
        self.assertEqual(exported["readiness_status"], "analysis_required")
        self.assertTrue(exported["development_only"])
        self.assertIn("source_consulted", markdown)
        self.assertNotIn("chain_of_thought", json.dumps(exported))

    def test_report_listing_and_csv_export_are_stable_and_utf8(self):
        report = self.open()
        self.store.append_event(
            report["report_id"], event_key="turn-4-completed",
            event_type="step_completed", step_id="candidate-a-step-04",
            step_label="4", user_wording="현재 단계를 완료했어요.",
            payload={
                "completion_source": "user_command",
                "pre_transition_step_id": "candidate-a-step-04",
                "post_transition_step_id": "candidate-a-step-05",
            },
        )
        listed = self.store.list_reports(session_id="session-test-1")
        self.assertEqual([item["report_id"] for item in listed], [report["report_id"]])
        content = self.store.export_csv(report["report_id"])
        self.assertTrue(content.startswith(b"\xef\xbb\xbf"))
        rows = list(csv.DictReader(io.StringIO(content.decode("utf-8-sig"))))
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["event_type"], "step_completed")
        self.assertEqual(rows[0]["step_label"], "4")
        self.assertEqual(rows[0]["user_wording"], "현재 단계를 완료했어요.")
        self.assertNotIn("chain_of_thought", content.decode("utf-8-sig"))

    def test_admin_aggregate_metrics_exclude_private_report_content(self):
        report = self.open()
        self.store.append_event(
            report["report_id"], event_key="timer-1", event_type="timer_started",
            step_id="candidate-a-step-03", step_label="3",
            user_wording="private timer wording",
        )
        self.store.append_event(
            report["report_id"], event_key="block-1", event_type="blocked",
            step_id="candidate-a-step-03", step_label="3",
            user_wording="private blocker wording",
        )
        metrics = self.store.aggregate_metrics()
        self.assertEqual(metrics["reports"]["total"], 1)
        self.assertEqual(metrics["workflow_events"]["timer_started"], 1)
        self.assertEqual(metrics["quality"]["blockers"], 1)
        self.assertEqual(
            metrics["quality"]["common_blocked_steps"],
            [{"step_label": "3", "count": 1}],
        )
        encoded = json.dumps(metrics)
        self.assertNotIn(report["report_id"], encoded)
        self.assertNotIn("session-test-1", encoded)
        self.assertNotIn("private", encoded)
        self.assertFalse(metrics["privacy"]["raw_audio_included"])

    def test_docx_export_uses_human_readable_labels_and_blank_student_fields(self):
        report = self.open()
        self.store.append_event(
            report["report_id"], event_key="turn-3-completed",
            event_type="step_completed", step_id="candidate-a-step-03",
            step_label="3", user_wording="현재 단계를 완료했어요.",
            payload={
                "timer": {
                    "source_duration_seconds": 900,
                    "elapsed_seconds": 20,
                    "remaining_seconds": 880,
                },
            },
        )
        content = self.store.export_docx(report["report_id"])
        self.assertGreater(len(content), 100)
        self.assertTrue(content.startswith(b"PK"))
        from docx import Document
        document = Document(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Title:", text)
        self.assertIn("Course:", text)
        self.assertIn("Student number:", text)
        self.assertIn("Name:", text)
        self.assertIn("Advisor:", text)
        self.assertIn("I. Purpose", text)
        self.assertIn("II. Materials and Methods", text)
        self.assertIn("III. Results", text)
        self.assertIn("IV. Discussion", text)
        self.assertIn("V. Conclusion", text)
        self.assertIn("Step 3 완료", text)
        self.assertIn("타이머 총 15:00", text)
        self.assertIn("경과 00:20", text)
        self.assertIn("잔여 14:40", text)
        self.assertNotIn("Student number: 20", text)
        self.assertNotIn("홍길동", text)
        self.assertNotIn("chain_of_thought", text)

    def test_grounded_report_context_and_scientific_narrative_integrity(self):
        from voice_workflow_agent.experiment_reports import (
            StepExecutionContext,
            GroundedReportContext,
            ReportWriterBrain,
            ReportWriterSettings,
            build_grounded_report_context,
        )
        report = self.open()
        # Record step completion with execution snapshot containing instruction, expected result, and actual observation
        self.store.append_event(
            report["report_id"],
            event_key="turn-1-complete",
            event_type="step_completed",
            step_id="candidate-a-step-01",
            step_label="1",
            user_wording="젤 밴드를 1x1 mm 크기로 잘랐습니다.",
            payload={
                "step_snapshot": {
                    "step_id": "candidate-a-step-01",
                    "step_label": "1",
                    "instruction": "Cut band into 1 mm3 pieces.",
                    "expected_results": ["1x1 mm pieces without contamination"],
                    "warnings": ["Clean surfaces, use scalpel and gloves."],
                    "source_page": 1,
                },
                "observations": ["투명하고 균일한 1 mm 크기로 절단 완료"],
            },
        )
        # Record second step where no actual observation was recorded
        self.store.append_event(
            report["report_id"],
            event_key="turn-2-complete",
            event_type="step_completed",
            step_id="candidate-a-step-02",
            step_label="2",
            user_wording="워시 완충액을 넣었습니다.",
            payload={
                "step_snapshot": {
                    "step_id": "candidate-a-step-02",
                    "step_label": "2",
                    "instruction": "Wash with 100 uL of 100 mM ABC/ACN (1:1, vol/vol).",
                    "expected_results": ["Gel destained completely"],
                    "source_page": 1,
                },
                "observations": [],
            },
        )
        doc = self.store.get_report(report["report_id"])
        ctx = build_grounded_report_context(doc)
        self.assertEqual(len(ctx.executed_steps), 2)
        self.assertEqual(ctx.executed_steps[0].step_label, "1")
        self.assertIn("Cut band into 1 mm3 pieces", ctx.executed_steps[0].instruction)
        self.assertEqual(ctx.executed_steps[0].user_confirmed_observations, ("투명하고 균일한 1 mm 크기로 절단 완료",))
        self.assertEqual(ctx.executed_steps[1].user_confirmed_observations, ())

        brain = ReportWriterBrain(ReportWriterSettings(enabled=True, model="grok-4.6-preview"))
        narrative = brain.build_deterministic_narrative(ctx)
        self.assertIn("목적", narrative.objective)
        self.assertIn("절차", narrative.materials_and_methods)
        # Results should mention actual observation for step 1
        self.assertIn("투명하고 균일한 1 mm 크기", narrative.results_and_observations)
        # Results should strictly disclaim unobserved actual result for step 2 without claiming false success
        self.assertIn("별도의 실제 관찰값은 기록되지 않았다", narrative.results_and_observations)
        self.assertNotIn("모든 단계가 정상적으로 완료되었다", narrative.results_and_observations)

    def test_docx_export_contains_grounded_step_table(self):
        report = self.open()
        self.store.append_event(
            report["report_id"],
            event_key="turn-1-complete",
            event_type="step_completed",
            step_id="step-1",
            step_label="1",
            user_wording="시약을 추가했습니다.",
            payload={
                "step_snapshot": {
                    "step_id": "step-1",
                    "step_label": "1",
                    "instruction": "Add 50 uL ABC buffer.",
                    "source_page": 1,
                },
            },
        )
        docx_bytes = self.store.export_docx(report["report_id"])
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        # Verify table exists in docx
        self.assertGreater(len(doc.tables), 0)
        table_text = " ".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
        self.assertIn("Step", table_text)
        self.assertIn("Event / Action", table_text)
        self.assertIn("시약을 추가했습니다", table_text)

    def test_grounded_report_with_candidate_a_domain_protocol_rehydration(self):
        """Test build_grounded_report_context with real domain ExperimentProtocol instance."""
        from voice_workflow_agent.curated_protocol import load_curated_protocol_fixture
        from voice_workflow_agent.experiment_reports import build_grounded_report_context, ReportWriterBrain

        fixture_path = Path(__file__).resolve().parents[1] / "data/development_protocols/candidate_a_curated_analysis.json"
        provenance_path = Path(__file__).resolve().parents[1] / "data/development_protocols/candidate_a_curated_analysis.provenance.json"
        pdf_path = (Path(__file__).resolve().parents[1] / "data" / "runtime" / "candidate-a-source" / "in-gel-digestion.pdf")
        fixture = load_curated_protocol_fixture(fixture_path, provenance_path, pdf_path)
        protocol = fixture.draft.protocol

        report = self.open()
        self.store.append_event(
            report["report_id"],
            event_key="turn-1-complete",
            event_type="step_completed",
            step_id="candidate-a-step-01",
            step_label="1",
            user_wording="젤 밴드를 절단했습니다.",
        )
        doc = self.store.get_report(report["report_id"])

        # Patch candidate fixture lookup to return the real candidate fixture
        with unittest.mock.patch("voice_workflow_agent.server._configured_candidate_fixture", return_value=fixture):
            ctx = build_grounded_report_context(doc)
            self.assertIsNotNone(ctx)
            self.assertEqual(ctx.protocol_id, "candidate-a-curated-development-v1")
            self.assertGreater(len(ctx.all_steps), 0)
            self.assertGreater(len(ctx.executed_steps), 0)
            self.assertEqual(ctx.executed_steps[0].step_label, "1")
            self.assertTrue(len(ctx.executed_steps[0].instruction_source_text) > 0)

            # Export docx
            docx_bytes = self.store.export_docx(report["report_id"])
            self.assertTrue(docx_bytes.startswith(b"PK"))
            from docx import Document
            docx_doc = Document(io.BytesIO(docx_bytes))
            self.assertGreater(len(docx_doc.paragraphs), 5)

    def test_http_docx_export_returns_200_and_valid_document(self):
        """Direct HTTP test for GET /api/experiment-reports/{report_id}.docx."""
        import httpx
        from voice_workflow_agent.server import app

        async def run_inline(function, *args, **kwargs):
            return function(*args, **kwargs)

        async def fetch_report(url: str):
            transport = httpx.ASGITransport(app=app)
            async with httpx.AsyncClient(
                transport=transport, base_url="http://testserver"
            ) as client:
                return await client.get(url)

        report = self.open()
        self.store.append_event(
            report["report_id"],
            event_key="turn-1-complete",
            event_type="step_completed",
            step_id="step-1",
            step_label="1",
            user_wording="1단계 작업을 마쳤습니다.",
        )

        with unittest.mock.patch.dict(os.environ, {
            "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORTS_ENABLED": "true",
            "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORT_DB": str(self.store.path),
            "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORTS_DATABASE": str(self.store.path),
            "VOICE_WORKFLOW_AGENT_REPORT_WRITER_ENABLED": "false",
        }), unittest.mock.patch(
            "fastapi.routing.run_in_threadpool", side_effect=run_inline,
        ):
            response = asyncio.run(fetch_report(
                f"/api/experiment-reports/{report['report_id']}.docx"
            ))
            self.assertEqual(response.status_code, 200)
            self.assertEqual(
                response.headers["content-type"],
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self.assertTrue(response.content.startswith(b"PK"))
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            self.assertGreater(len(doc.paragraphs), 5)


if __name__ == "__main__":
    unittest.main()
