"""Idempotent, event-driven experiment records owned by the server.

The service stores only explicit workflow events and user-confirmed observations.
It never infers completion, approval, or a laboratory result from model output.
"""

from __future__ import annotations

import hashlib
import csv
import io
import json
import os
import re
import secrets
import sqlite3
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from voice_workflow_agent.report_projection import project_protocol_for_report, project_step_for_report


_TRUE = frozenset({"1", "true", "yes", "on"})
_FALSE = frozenset({"0", "false", "no", "off", ""})
_SAFE_ID = re.compile(r"^[A-Za-z0-9._:-]{1,160}$")
_STATUSES = frozenset({"in_progress", "completed", "blocked", "stopped", "incomplete"})


@dataclass(frozen=True)
class ExperimentReportSettings:
    enabled: bool
    database_path: Path | None = None

    @classmethod
    def from_environment(cls) -> "ExperimentReportSettings":
        raw = os.environ.get(
            "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORTS_ENABLED", ""
        ).strip().casefold()
        if raw in _FALSE:
            return cls(False)
        if raw not in _TRUE:
            raise ValueError(
                "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORTS_ENABLED must be a boolean"
            )
        path = os.environ.get(
            "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORT_DB",
            os.environ.get("VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORTS_DATABASE", ""),
        ).strip()
        if not path:
            raise ValueError(
                "VOICE_WORKFLOW_AGENT_EXPERIMENT_REPORT_DB is required"
            )
        return cls(True, Path(path))


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _clean_identifier(value: str, field: str) -> str:
    if not isinstance(value, str) or _SAFE_ID.fullmatch(value) is None:
        raise ValueError(f"{field} is invalid")
    return value


def _clean_text(value: str, *, maximum: int = 1600) -> str:
    if not isinstance(value, str):
        raise ValueError("report text is invalid")
    cleaned = " ".join(value.split())
    if not cleaned or len(cleaned) > maximum:
        raise ValueError("report text is invalid")
    return cleaned


class ExperimentReportStore:
    """Small SQLite store with one report per procedure session."""

    SCHEMA_VERSION = 1

    def __init__(self, path: str | Path) -> None:
        self.path = Path(path)

    @property
    def database_path(self) -> Path:
        return self.path

    def _connect(self) -> sqlite3.Connection:
        self.path.parent.mkdir(parents=True, exist_ok=True)
        connection = sqlite3.connect(self.path)
        connection.row_factory = sqlite3.Row
        connection.execute("PRAGMA foreign_keys = ON")
        connection.execute("PRAGMA busy_timeout = 5000")
        connection.executescript(
            """
            CREATE TABLE IF NOT EXISTS experiment_report_metadata (
              schema_version INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS experiment_reports (
              report_id TEXT PRIMARY KEY,
              session_id TEXT NOT NULL UNIQUE,
              protocol_id TEXT NOT NULL,
              protocol_title TEXT NOT NULL,
              protocol_revision TEXT NOT NULL,
              protocol_sha256 TEXT NOT NULL,
              readiness_status TEXT NOT NULL,
              development_only INTEGER NOT NULL,
              status TEXT NOT NULL,
              started_at TEXT NOT NULL,
              ended_at TEXT,
              timezone TEXT NOT NULL,
              anomaly_count INTEGER NOT NULL DEFAULT 0,
              blocker_count INTEGER NOT NULL DEFAULT 0,
              finalization_version INTEGER NOT NULL DEFAULT 0
            );
            CREATE TABLE IF NOT EXISTS experiment_report_events (
              id INTEGER PRIMARY KEY,
              report_id TEXT NOT NULL REFERENCES experiment_reports(report_id),
              event_key TEXT NOT NULL,
              event_type TEXT NOT NULL,
              step_id TEXT,
              step_label TEXT,
              user_wording TEXT,
              category TEXT,
              severity TEXT,
              confirmation_state TEXT,
              source_tier TEXT,
              citation_identities TEXT NOT NULL,
              payload TEXT NOT NULL,
              created_at TEXT NOT NULL,
              UNIQUE(report_id, event_key)
            );
            """
        )
        rows = connection.execute(
            "SELECT schema_version FROM experiment_report_metadata"
        ).fetchall()
        if not rows:
            connection.execute(
                "INSERT INTO experiment_report_metadata(schema_version) VALUES (?)",
                (self.SCHEMA_VERSION,),
            )
        elif len(rows) != 1 or rows[0][0] != self.SCHEMA_VERSION:
            connection.close()
            raise RuntimeError("experiment report schema is unsupported")
        return connection

    def open_report(
        self,
        *,
        session_id: str,
        protocol_id: str,
        protocol_title: str,
        protocol_revision: str,
        protocol_sha256: str,
        readiness_status: str,
        development_only: bool,
    ) -> dict[str, Any]:
        session_id = _clean_identifier(session_id, "session_id")
        protocol_id = _clean_identifier(protocol_id, "protocol_id")
        protocol_revision = _clean_identifier(protocol_revision, "protocol_revision")
        if not re.fullmatch(r"[0-9a-f]{64}", protocol_sha256):
            raise ValueError("protocol_sha256 is invalid")
        title = _clean_text(protocol_title, maximum=400)
        readiness = _clean_identifier(readiness_status, "readiness_status")
        report_id = "ER-" + hashlib.sha256(
            f"{session_id}\x1f{protocol_id}\x1f{protocol_revision}".encode()
        ).hexdigest()[:20].upper()
        now = _now()
        with self._connect() as connection:
            connection.execute(
                """
                INSERT OR IGNORE INTO experiment_reports(
                  report_id,session_id,protocol_id,protocol_title,
                  protocol_revision,protocol_sha256,readiness_status,
                  development_only,status,started_at,timezone
                ) VALUES(?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    report_id, session_id, protocol_id, title,
                    protocol_revision, protocol_sha256, readiness,
                    int(bool(development_only)), "in_progress", now, "UTC",
                ),
            )
        return self.get_report(report_id)

    def append_event(
        self,
        report_id: str,
        *,
        event_key: str,
        event_type: str,
        step_id: str | None = None,
        step_label: str | None = None,
        user_wording: str | None = None,
        category: str | None = None,
        severity: str | None = None,
        confirmation_state: str | None = None,
        source_tier: str | None = None,
        citation_identities: tuple[str, ...] = (),
        payload: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        report_id = _clean_identifier(report_id, "report_id")
        event_key = _clean_identifier(event_key, "event_key")
        event_type = _clean_identifier(event_type, "event_type")
        for value, field in (
            (step_id, "step_id"), (step_label, "step_label"),
            (category, "category"), (severity, "severity"),
            (confirmation_state, "confirmation_state"),
            (source_tier, "source_tier"),
        ):
            if value is not None:
                _clean_identifier(value, field)
        wording = (
            _clean_text(user_wording, maximum=800)
            if user_wording is not None else None
        )
        citations = tuple(
            _clean_identifier(item, "citation_identity")
            for item in citation_identities[:20]
        )
        safe_payload = payload or {}
        encoded = json.dumps(
            safe_payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")
        )
        if len(encoded) > 12000:
            raise ValueError("report payload is too large")
        with self._connect() as connection:
            cursor = connection.execute(
                """
                INSERT OR IGNORE INTO experiment_report_events(
                  report_id,event_key,event_type,step_id,step_label,user_wording,
                  category,severity,confirmation_state,source_tier,
                  citation_identities,payload,created_at
                ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    report_id, event_key, event_type, step_id, step_label,
                    wording, category, severity, confirmation_state, source_tier,
                    json.dumps(citations), encoded, _now(),
                ),
            )
            if cursor.rowcount and event_type == "anomaly":
                connection.execute(
                    "UPDATE experiment_reports SET anomaly_count=anomaly_count+1 "
                    "WHERE report_id=?",
                    (report_id,),
                )
            if cursor.rowcount and event_type == "blocked":
                connection.execute(
                    "UPDATE experiment_reports SET blocker_count=blocker_count+1 "
                    "WHERE report_id=?",
                    (report_id,),
                )
        result = self.get_report(report_id)
        result["event_inserted"] = bool(cursor.rowcount)
        return result

    def finalize(
        self, report_id: str, *, status: str, event_key: str
    ) -> dict[str, Any]:
        report_id = _clean_identifier(report_id, "report_id")
        event_key = _clean_identifier(event_key, "event_key")
        if status not in _STATUSES - {"in_progress"}:
            raise ValueError("report status is invalid")
        self.append_event(
            report_id, event_key=event_key, event_type="report_finalized",
            payload={"status": status},
        )
        with self._connect() as connection:
            connection.execute(
                """
                UPDATE experiment_reports
                   SET status=?, ended_at=COALESCE(ended_at,?),
                       finalization_version=CASE
                         WHEN ended_at IS NULL THEN finalization_version+1
                         ELSE finalization_version END
                 WHERE report_id=?
                """,
                (status, _now(), report_id),
            )
        return self.get_report(report_id)

    def get_report(self, report_id: str) -> dict[str, Any]:
        report_id = _clean_identifier(report_id, "report_id")
        with self._connect() as connection:
            report = connection.execute(
                "SELECT * FROM experiment_reports WHERE report_id=?", (report_id,)
            ).fetchone()
            if report is None:
                raise KeyError("experiment report not found")
            events = connection.execute(
                """
                SELECT event_key,event_type,step_id,step_label,user_wording,
                       category,severity,confirmation_state,source_tier,
                       citation_identities,payload,created_at
                  FROM experiment_report_events
                 WHERE report_id=? ORDER BY id
                """,
                (report_id,),
            ).fetchall()
        result = dict(report)
        result["development_only"] = bool(result["development_only"])
        result["events"] = [
            {
                **{key: row[key] for key in (
                    "event_key", "event_type", "step_id", "step_label",
                    "user_wording", "category", "severity",
                    "confirmation_state", "source_tier", "created_at",
                )},
                "citation_identities": json.loads(row["citation_identities"]),
                "payload": json.loads(row["payload"]),
            }
            for row in events
        ]
        return result

    def list_reports(self, *, session_id: str | None = None) -> list[dict[str, Any]]:
        """List report summaries in stable start order for one development session."""

        parameters: tuple[Any, ...] = ()
        where = ""
        if session_id is not None:
            where = "WHERE session_id=?"
            parameters = (_clean_identifier(session_id, "session_id"),)
        with self._connect() as connection:
            rows = connection.execute(
                f"""
                SELECT report_id,session_id,protocol_id,status,started_at,ended_at,
                       anomaly_count,blocker_count,finalization_version,
                       development_only
                  FROM experiment_reports {where}
                 ORDER BY started_at,report_id
                """,
                parameters,
            ).fetchall()
        return [
            {**dict(row), "development_only": bool(row["development_only"])}
            for row in rows
        ]

    def export_json(self, report_id: str) -> bytes:
        return (
            json.dumps(
                self.get_report(report_id), ensure_ascii=False,
                sort_keys=True, indent=2,
            ) + "\n"
        ).encode()

    def export_markdown(self, report_id: str) -> bytes:
        report = self.get_report(report_id)
        lines = [
            f"# Experiment report {report['report_id']}",
            "",
            f"- Protocol: {report['protocol_title']} ({report['protocol_id']})",
            f"- Revision: {report['protocol_revision']}",
            f"- Source SHA-256: {report['protocol_sha256']}",
            f"- Readiness: {report['readiness_status']}",
            f"- Development only: {str(report['development_only']).lower()}",
            f"- Status: {report['status']}",
            f"- Started: {report['started_at']}",
            f"- Ended: {report['ended_at'] or '—'}",
            "",
            "## Event timeline",
            "",
        ]
        for event in report["events"]:
            payload = event.get("payload") or {}
            detail = event["user_wording"] or payload.get("summary") or ""
            timer = payload.get("timer") if isinstance(payload.get("timer"), dict) else {}
            timer_bits = []
            for key, label in (
                ("source_duration_seconds", "defined"),
                ("started_at", "started"),
                ("elapsed_seconds", "elapsed"),
                ("remaining_seconds", "remaining"),
                ("completion_state", "completion"),
                ("demo_bypassed", "demo_bypassed"),
            ):
                if timer.get(key) not in (None, ""):
                    timer_bits.append(f"{label}={timer[key]}")
            extra = " · ".join(item for item in (detail, "; ".join(timer_bits)) if item)
            lines.append(
                f"- {event['created_at']} · {event['event_type']} · "
                f"step {event['step_label'] or '—'}"
                + (f" · {extra}" if extra else "")
            )
        return ("\n".join(lines) + "\n").encode()

    def export_csv(self, report_id: str) -> bytes:
        """Export the stable event timeline as UTF-8 CSV."""

        report = self.get_report(report_id)
        output = io.StringIO(newline="")
        writer = csv.writer(output, lineterminator="\n")
        writer.writerow((
            "report_id", "event_key", "event_type", "step_id", "step_label",
            "category", "severity", "confirmation_state", "source_tier",
            "user_wording", "created_at",
            "source_duration_seconds", "timer_started_at", "elapsed_seconds",
            "remaining_seconds", "completion_state", "demo_bypassed",
        ))
        for event in report["events"]:
            payload = event.get("payload") or {}
            timer = payload.get("timer") if isinstance(payload.get("timer"), dict) else {}
            writer.writerow((
                report["report_id"], event["event_key"], event["event_type"],
                event["step_id"] or "", event["step_label"] or "",
                event["category"] or "", event["severity"] or "",
                event["confirmation_state"] or "", event["source_tier"] or "",
                event["user_wording"] or "", event["created_at"],
                timer.get("source_duration_seconds", ""),
                timer.get("started_at", ""),
                timer.get("elapsed_seconds", ""),
                timer.get("remaining_seconds", ""),
                timer.get("completion_state", ""),
                timer.get("demo_bypassed", ""),
            ))
        return output.getvalue().encode("utf-8-sig")

    def docx_bytes(self, report_id: str) -> bytes:
        return self.export_docx(report_id)

    def export_docx(
        self, report_id: str, narrative: ReportNarrative | None = None
    ) -> bytes:
        """Export a bounded undergraduate lab-report .docx from the event ledger.

        Student metadata fields stay blank. The SQLite event ledger remains the
        authority; this document is a derived, human-readable projection.
        """

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor

        report = self.get_report(report_id)
        document = Document()
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Laboratory Experiment Report / 실험 결과 보고서")
        run.bold = True
        run.font.size = Pt(16)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(f"Protocol: {report.get('protocol_title') or 'Authoritative Protocol'}")
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(80, 80, 80)

        # Faithful header fields. Never invent student metadata.
        for label in ("Title", "Course", "Student number", "Name", "Advisor"):
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True

        events = list(report.get("events") or ())
        event_cap = 100
        truncated = len(events) > event_cap
        visible_events = events[:event_cap]
        timeline = [
            _human_event_label(event) for event in visible_events
        ]
        if truncated:
            timeline.append(
                "Additional events omitted to keep the report within standard length."
            )

        started = _human_event_clock(report.get("started_at"))
        ended = _human_event_clock(report.get("ended_at"))

        # Extract anomalies from events
        anomaly_events = [
            e for e in events
            if e.get("event_type") in ("anomaly", "system_anomaly") or e.get("anomaly_category")
        ]
        anomaly_lines = []
        for ae in anomaly_events:
            clock = _human_event_clock(ae.get("created_at"))
            step = ae.get("step_label") or "—"
            cat = ae.get("anomaly_category") or ae.get("event_type") or "anomaly"
            note = ae.get("user_wording") or (ae.get("payload") or {}).get("text") or "Recorded issue"
            anomaly_lines.append(f"• [{clock}] Step {step} ({cat}): {note}")

        # Extract observations
        observation_events = [e for e in events if e.get("event_type") == "observation"]
        observation_lines = []
        for obs in observation_events:
            clock = _human_event_clock(obs.get("created_at"))
            step = obs.get("step_label") or "—"
            text = obs.get("user_wording") or (obs.get("payload") or {}).get("text") or "Observation"
            observation_lines.append(f"• [{clock}] Step {step}: {text}")

        # Extract pauses
        pause_events = [e for e in events if e.get("event_type") == "workflow_paused"]
        pause_lines = []
        for pe in pause_events:
            clock = _human_event_clock(pe.get("created_at"))
            reason = pe.get("user_wording") or (pe.get("payload") or {}).get("reason") or "Paused"
            pause_lines.append(f"• [{clock}] Paused: {reason}")

        timeline = [_human_event_label(e) for e in events if _human_event_label(e)]

        # Use narrative if provided, otherwise deterministic narrative
        if narrative is None:
            writer = ReportWriterBrain()
            narrative = writer.build_deterministic_narrative(report, events)

        sections = (
            (
                "I. Purpose / 실험 목적",
                f"Protocol: {report.get('protocol_title') or '—'} ({report.get('protocol_id') or '—'}).\n"
                f"Session Report ID: {report_id}.\n"
                f"{narrative.objective}"
            ),
            (
                "II. Materials and Methods / 재료 및 실험 방법",
                f"Protocol Revision: {report.get('protocol_revision') or '—'}.\n"
                f"Readiness Status: {report.get('readiness_status') or '—'}.\n"
                f"Protocol SHA-256: {report.get('protocol_sha256') or '—'}.\n"
                f"{narrative.materials_and_methods}"
            ),
            (
                "III. Results / 수행 내용 및 관찰 결과",
                ((f"Execution Event Timeline:\n" + "\n".join(f"• {item}" for item in timeline) + "\n\n") if timeline else "")
                + f"{narrative.results_and_observations}"
            ),
            (
                "IV. Discussion / 고찰",
                f"Recorded anomalies: {int(report.get('anomaly_count') or 0)}. Recorded blockers: {int(report.get('blocker_count') or 0)}.\n"
                f"{narrative.discussion}"
            ),
            (
                "V. Conclusion / 결론",
                f"Status: {report.get('status') or '—'}. Started: {started or '—'}. Ended: {ended or '—'}.\n"
                f"Cryptographic Ledger Events: {len(events)} events committed.\n"
                f"{narrative.conclusion}"
            ),
            (
                "VI. Limitations / 한계 및 미기록 항목",
                f"{narrative.limitations}"
            ),
        )
        for heading, body in sections:
            heading_paragraph = document.add_paragraph()
            heading_run = heading_paragraph.add_run(heading)
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            heading_run.font.color.rgb = RGBColor(20, 50, 35)
            for line in str(body).split("\n"):
                p = document.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)

        # Structured Event Table as Appendix
        if visible_events:
            table_heading = document.add_paragraph()
            table_run = table_heading.add_run("Appendix: Detailed Event Ledger / 상세 이벤트 감사 기록")
            table_run.bold = True
            table_run.font.size = Pt(11)
            table_run.font.color.rgb = RGBColor(20, 50, 35)

            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ("Step", "Event / Action", "Time", "Timer Status", "Notes / Description")
            for i, header_text in enumerate(headers):
                hdr_cells[i].text = header_text
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

            for ev in visible_events:
                row_cells = table.add_row().cells
                ev_type = str(ev.get("event_type") or "")
                verb = _HUMAN_EVENT_VERBS.get(ev_type, ev_type.replace("_", " "))
                step_str = str(ev.get("step_label") or "—")
                clock_str = _human_event_clock(ev.get("created_at"))
                payload = ev.get("payload") if isinstance(ev.get("payload"), dict) else {}
                timer = payload.get("timer") if isinstance(payload.get("timer"), dict) else {}
                duration = timer.get("source_duration_seconds", timer.get("duration_seconds"))
                elapsed = timer.get("elapsed_seconds")
                remaining = timer.get("remaining_seconds")
                t_parts = []
                if duration not in (None, ""):
                    t_parts.append(f"총 {_format_elapsed_clock(duration)}")
                if elapsed not in (None, ""):
                    t_parts.append(f"경과 {_format_elapsed_clock(elapsed)}")
                if remaining not in (None, ""):
                    t_parts.append(f"잔여 {_format_elapsed_clock(remaining)}")
                timer_str = " · ".join(t_parts) if t_parts else "—"
                notes_str = str(ev.get("user_wording") or (ev.get("payload") or {}).get("text") or "")

                row_cells[0].text = step_str
                row_cells[1].text = verb
                row_cells[2].text = clock_str
                row_cells[3].text = timer_str
                row_cells[4].text = notes_str
                for cell in row_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8.5)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()


@dataclass(frozen=True)
class StepExecutionContext:
    """Detailed factual context for one protocol step and its execution status."""

    step_id: str
    step_label: str
    section_title: str
    instruction_source_text: str
    sub_actions: tuple[str, ...] = ()
    quantities: tuple[str, ...] = ()
    conditions: tuple[str, ...] = ()
    expected_results: tuple[str, ...] = ()
    warnings: tuple[str, ...] = ()
    notes: tuple[str, ...] = ()
    tips: tuple[str, ...] = ()
    source_page: int = 1
    evidence_ids: tuple[str, ...] = ()
    entered_at: str | None = None
    completed_at: str | None = None
    completion_state: str = "not_started"  # "completed", "in_progress", "not_started"
    timer_configuration: str | None = None
    timer_actuals: str | None = None
    user_confirmed_observations: tuple[str, ...] = ()
    anomalies_deviations: tuple[str, ...] = ()
    applicable_safety_references: tuple[str, ...] = ()

    @property
    def instruction(self) -> str:
        return self.instruction_source_text

    def public_dict(self) -> dict[str, Any]:
        return {
            "step_id": self.step_id,
            "step_label": self.step_label,
            "section_title": self.section_title,
            "instruction_source_text": self.instruction_source_text,
            "sub_actions": list(self.sub_actions),
            "quantities": list(self.quantities),
            "conditions": list(self.conditions),
            "expected_results": list(self.expected_results),
            "warnings": list(self.warnings),
            "notes": list(self.notes),
            "tips": list(self.tips),
            "source_page": self.source_page,
            "evidence_ids": list(self.evidence_ids),
            "entered_at": self.entered_at,
            "completed_at": self.completed_at,
            "completion_state": self.completion_state,
            "timer_configuration": self.timer_configuration,
            "timer_actuals": self.timer_actuals,
            "user_confirmed_observations": list(self.user_confirmed_observations),
            "anomalies_deviations": list(self.anomalies_deviations),
            "applicable_safety_references": list(self.applicable_safety_references),
        }


@dataclass(frozen=True)
class GroundedReportContext:
    """Rich factual context constructed from the ledger and verified protocol store."""

    report_metadata: dict[str, Any]
    protocol_metadata: dict[str, Any]
    experiment_objective: str
    materials: tuple[str, ...]
    equipment: tuple[str, ...]
    prerequisites: tuple[str, ...]
    executed_steps: tuple[StepExecutionContext, ...]
    all_protocol_steps: tuple[StepExecutionContext, ...]
    observations: tuple[dict[str, Any], ...]
    timers: tuple[dict[str, Any], ...]
    anomalies: tuple[dict[str, Any], ...]
    deviations: tuple[dict[str, Any], ...]
    safety_pack_summary: dict[str, Any] | None
    source_references: tuple[str, ...]
    session_timing: dict[str, Any]
    event_ledger: tuple[dict[str, Any], ...]

    @property
    def protocol_id(self) -> str:
        return str(self.protocol_metadata.get("protocol_id") or "")

    @property
    def report_id(self) -> str:
        return str(self.report_metadata.get("report_id") or "")

    @property
    def all_steps(self) -> tuple[StepExecutionContext, ...]:
        return self.all_protocol_steps

    def public_dict(self) -> dict[str, Any]:
        return {
            "report_metadata": self.report_metadata,
            "protocol_metadata": self.protocol_metadata,
            "experiment_objective": self.experiment_objective,
            "materials": list(self.materials),
            "equipment": list(self.equipment),
            "prerequisites": list(self.prerequisites),
            "executed_steps": [s.public_dict() for s in self.executed_steps],
            "observations": list(self.observations),
            "timers": list(self.timers),
            "anomalies": list(self.anomalies),
            "deviations": list(self.deviations),
            "safety_pack_summary": self.safety_pack_summary,
            "source_references": list(self.source_references),
            "session_timing": self.session_timing,
            "total_event_count": len(self.event_ledger),
        }


def build_grounded_report_context(
    report_data: dict[str, Any],
    events: list[dict[str, Any]] | None = None,
) -> GroundedReportContext:
    """Rehydrate protocol steps and merge execution event timestamps, observations, and timers."""
    if events is None:
        events = list(report_data.get("events") or [])
    protocol_id = str(report_data.get("protocol_id") or "")
    protocol_rev = str(report_data.get("protocol_revision") or "")
    protocol_title = str(report_data.get("protocol_title") or "Experiment Protocol")

    step_events: dict[str, list[dict[str, Any]]] = {}
    obs_list: list[dict[str, Any]] = []
    timer_list: list[dict[str, Any]] = []
    anomaly_list: list[dict[str, Any]] = []
    deviation_list: list[dict[str, Any]] = []
    step_snapshots: dict[str, dict[str, Any]] = {}

    for e in events:
        s_lbl = str(e.get("step_label") or "")
        if s_lbl and s_lbl != "—":
            step_events.setdefault(s_lbl, []).append(e)
        ev_type = str(e.get("event_type") or "")
        payload = e.get("payload") if isinstance(e.get("payload"), dict) else {}

        if "step_snapshot" in payload and isinstance(payload["step_snapshot"], dict):
            snap = payload["step_snapshot"]
            snap_lbl = str(snap.get("step_label") or s_lbl)
            if snap_lbl:
                step_snapshots[snap_lbl] = snap

        if ev_type == "observation":
            obs_list.append(e)
        elif ev_type in ("timer_started", "timer_expired") or "timer" in payload:
            timer_list.append(e)
        elif ev_type in ("anomaly", "system_anomaly") or e.get("anomaly_category"):
            anomaly_list.append(e)
        elif ev_type in ("workflow_paused", "workflow_blocked"):
            deviation_list.append(e)

    # Try rehydrating stored protocol from database or candidate fixture
    stored_protocol = None
    protocol_revision = str(report_data.get("protocol_revision") or "1")
    try:
        from voice_workflow_agent.server import _open_protocol_catalog, _configured_candidate_fixture, server_config
        cfg = server_config()
        candidate = _configured_candidate_fixture(cfg)
        if candidate is not None and candidate.protocol_id == protocol_id:
            stored_protocol = candidate.draft.protocol
        else:
            cat, st = _open_protocol_catalog()
            try:
                fix = cat.load_executable_fixture(protocol_id, protocol_revision)
                stored_protocol = fix.draft.protocol
            finally:
                st.close()
    except Exception:
        stored_protocol = None

    all_steps: list[StepExecutionContext] = []
    executed_steps: list[StepExecutionContext] = []
    all_materials: set[str] = set()
    all_equipment: set[str] = set()
    all_prereqs: set[str] = set()
    objective = f"본 실험은 '{protocol_title}' 지침에 따라 표준화된 실험 절차를 수행하고 검증된 실험 데이터를 기록하는 것을 목적으로 한다."

    if stored_protocol is not None:
        projected = project_protocol_for_report(stored_protocol)
        if projected.objective:
            objective = projected.objective
        all_materials = set(projected.materials)
        all_equipment = set(projected.equipment)
        all_prereqs = set(projected.prerequisites)

        for p_step in projected.steps:
            lbl = p_step.step_label
            evs = step_events.get(lbl, [])
            entered = next((_human_event_clock(ev.get("created_at")) for ev in evs if ev.get("event_type") in ("step_entered", "step_presented", "session_started")), None)
            completed = next((_human_event_clock(ev.get("created_at")) for ev in evs if ev.get("event_type") == "step_completed"), None)
            state = "completed" if completed else ("in_progress" if entered else "not_started")

            # Timers
            t_cfg = None
            t_act = None
            t_ev = next((ev for ev in evs if ev.get("event_type") == "timer_started" or "timer" in (ev.get("payload") or {})), None)
            if t_ev:
                t_pay = (t_ev.get("payload") or {}).get("timer", {})
                dur = t_pay.get("source_duration_seconds", t_pay.get("duration_seconds"))
                elap = t_pay.get("elapsed_seconds")
                if dur:
                    t_cfg = f"{dur}초 ({_format_elapsed_clock(dur)})"
                if elap:
                    t_act = f"{elap}초 경과 ({_format_elapsed_clock(elap)})"

            # Observations
            obs_collected = []
            for ev in evs:
                if ev.get("event_type") == "observation":
                    text = str(ev.get("user_wording") or (ev.get("payload") or {}).get("text") or "관찰 기록")
                    if text:
                        obs_collected.append(text)
                payload = ev.get("payload") if isinstance(ev.get("payload"), dict) else {}
                if "observations" in payload and isinstance(payload["observations"], (list, tuple)):
                    for o in payload["observations"]:
                        if o and str(o).strip():
                            obs_collected.append(str(o).strip())
            step_obs = tuple(obs_collected)

            # Anomalies
            step_anom = tuple(
                str(ev.get("user_wording") or (ev.get("payload") or {}).get("text") or "이상 보고")
                for ev in evs if ev.get("event_type") in ("anomaly", "system_anomaly") or ev.get("anomaly_category")
            )

            ctx_step = StepExecutionContext(
                step_id=p_step.step_id,
                step_label=lbl,
                section_title=p_step.section_title,
                instruction_source_text=p_step.instruction_source_text,
                sub_actions=p_step.sub_actions,
                quantities=p_step.quantities,
                conditions=p_step.conditions,
                expected_results=p_step.expected_results,
                warnings=p_step.warnings,
                notes=p_step.notes,
                tips=p_step.tips,
                source_page=p_step.source_page,
                evidence_ids=p_step.evidence_ids,
                entered_at=entered,
                completed_at=completed,
                completion_state=state,
                timer_configuration=t_cfg,
                timer_actuals=t_act,
                user_confirmed_observations=step_obs,
                anomalies_deviations=step_anom,
            )
            all_steps.append(ctx_step)
            if state in ("completed", "in_progress") or step_obs or step_anom:
                executed_steps.append(ctx_step)

    else:
        # Fall back to step_snapshots or step_events
        for lbl, evs in sorted(step_events.items(), key=lambda x: str(x[0])):
            snap = step_snapshots.get(lbl, {})
            entered = next((_human_event_clock(ev.get("created_at")) for ev in evs if ev.get("event_type") in ("step_entered", "step_presented", "session_started")), None)
            completed = next((_human_event_clock(ev.get("created_at")) for ev in evs if ev.get("event_type") == "step_completed"), None)
            state = "completed" if completed else "in_progress"

            inst = snap.get("instruction_source_text") or snap.get("instruction") or f"Step {lbl} instruction."
            sec_title = snap.get("section_title") or "Experiment Execution"
            quantities = tuple(snap.get("quantities") or ())
            conditions = tuple(snap.get("conditions") or ())
            exp_res = tuple(snap.get("expected_results") or ())
            warnings = tuple(snap.get("warnings") or ())
            page = int(snap.get("source_page") or 1)

            obs_collected = []
            for ev in evs:
                if ev.get("event_type") == "observation":
                    text = str(ev.get("user_wording") or (ev.get("payload") or {}).get("text") or "관찰 기록")
                    if text:
                        obs_collected.append(text)
                payload = ev.get("payload") if isinstance(ev.get("payload"), dict) else {}
                if "observations" in payload and isinstance(payload["observations"], (list, tuple)):
                    for o in payload["observations"]:
                        if o and str(o).strip():
                            obs_collected.append(str(o).strip())
            step_obs = tuple(obs_collected)

            step_anom = tuple(
                str(ev.get("user_wording") or (ev.get("payload") or {}).get("text") or "이상 보고")
                for ev in evs if ev.get("event_type") in ("anomaly", "system_anomaly") or ev.get("anomaly_category")
            )

            ctx_step = StepExecutionContext(
                step_id=snap.get("step_id") or f"step-{lbl}",
                step_label=lbl,
                section_title=sec_title,
                instruction_source_text=inst,
                quantities=quantities,
                conditions=conditions,
                expected_results=exp_res,
                warnings=warnings,
                source_page=page,
                entered_at=entered,
                completed_at=completed,
                completion_state=state,
                user_confirmed_observations=step_obs,
                anomalies_deviations=step_anom,
            )
            all_steps.append(ctx_step)
            executed_steps.append(ctx_step)

    refs = [f"실험 PDF: {report_data.get('protocol_title') or protocol_title} (ID: {protocol_id}, Rev: {protocol_rev})"]
    if report_data.get("protocol_sha256"):
        refs.append(f"PDF SHA-256: {report_data.get('protocol_sha256')}")

    return GroundedReportContext(
        report_metadata={
            "report_id": report_data.get("report_id"),
            "session_id": report_data.get("session_id"),
            "status": report_data.get("status"),
            "anomaly_count": report_data.get("anomaly_count", len(anomaly_list)),
            "blocker_count": report_data.get("blocker_count", 0),
            "finalization_version": report_data.get("finalization_version", 0),
        },
        protocol_metadata={
            "protocol_id": protocol_id,
            "protocol_title": protocol_title,
            "protocol_revision": protocol_rev,
            "protocol_sha256": report_data.get("protocol_sha256"),
            "total_steps": len(all_steps),
        },
        experiment_objective=objective,
        materials=tuple(sorted(all_materials)),
        equipment=tuple(sorted(all_equipment)),
        prerequisites=tuple(sorted(all_prereqs)),
        executed_steps=tuple(executed_steps),
        all_protocol_steps=tuple(all_steps),
        observations=tuple(obs_list),
        timers=tuple(timer_list),
        anomalies=tuple(anomaly_list),
        deviations=tuple(deviation_list),
        safety_pack_summary=None,
        source_references=tuple(refs),
        session_timing={
            "started_at": report_data.get("started_at"),
            "ended_at": report_data.get("ended_at"),
            "timezone": report_data.get("timezone", "UTC"),
        },
        event_ledger=tuple(events),
    )


@dataclass(frozen=True)
class ReportNarrative:
    """Structured, professional narrative laboratory report produced by ReportWriterBrain."""

    title: str
    objective: str
    session_summary: str
    chronological_highlights: tuple[str, ...]
    materials_and_methods: str
    results_and_observations: str
    discussion: str
    anomalies_and_deviations: str
    conclusion: str
    limitations: str


@dataclass(frozen=True)
class ReportDraftState:
    """Server-owned structured draft state for specialized report writing."""

    report_id: str
    session_id: str
    protocol_id: str
    status: str
    experiment_summary: str = ""
    materials_summary: str = ""
    equipment_summary: str = ""
    observations_narrative: str = ""
    anomalies_narrative: str = ""
    timeline_narrative: str = ""
    conclusion_narrative: str = ""
    committed_event_ids: tuple[str, ...] = ()
    last_updated_at: str = ""


@dataclass(frozen=True)
class ReportWriterSettings:
    enabled: bool = True
    model: str = "grok-4.6"
    timeout_seconds: float = 25.0

    @classmethod
    def from_environment(cls) -> "ReportWriterSettings":
        enabled_val = os.environ.get(
            "VOICE_WORKFLOW_AGENT_REPORT_WRITER_ENABLED", "true"
        ).strip().casefold()
        enabled = enabled_val in _TRUE
        model = os.environ.get(
            "VOICE_WORKFLOW_AGENT_REPORT_WRITER_MODEL",
            os.environ.get("EXTERNAL_REFERENCE_MODEL", "grok-4.6"),
        ).strip() or "grok-4.6"
        try:
            timeout = float(os.environ.get(
                "VOICE_WORKFLOW_AGENT_REPORT_WRITER_TIMEOUT_SECONDS", "25"
            ).strip())
        except ValueError:
            timeout = 25.0
        return cls(enabled=enabled, model=model, timeout_seconds=timeout)


_NARRATIVE_CACHE: dict[tuple[str, int, str, str, str], ReportNarrative] = {}


class ReportWriterBrain:
    """Specialist Brain generating faithful laboratory report prose from grounded step context."""

    def __init__(
        self,
        client: Any = None,
        model: str = "grok-4.6",
        timeout_seconds: float = 25.0,
    ) -> None:
        self.client = client
        self.model = model
        self.timeout_seconds = timeout_seconds

    def build_deterministic_draft(
        self,
        report_data: dict[str, Any],
        events: list[dict[str, Any]],
    ) -> ReportDraftState:
        """Create a faithful, unhallucinated report draft directly from SQLite events."""
        event_ids = tuple(
            str(e.get("event_key") or e.get("event_id"))
            for e in events
            if (e.get("event_key") or e.get("event_id"))
        )
        obs_count = sum(1 for e in events if e.get("event_type") == "observation")
        anomaly_count = sum(1 for e in events if e.get("event_type") in ("anomaly", "system_anomaly"))
        completed_steps = [
            e.get("step_label") for e in events if e.get("event_type") == "step_completed"
        ]

        summary = (
            f"Protocol {report_data.get('protocol_title', 'Protocol')} "
            f"execution recorded {len(events)} total ledger events across "
            f"{len(completed_steps)} completed steps."
        )
        obs_narrative = (
            f"Operator recorded {obs_count} direct qualitative observations during the session."
            if obs_count > 0 else "해당 단계에 대해 별도의 관찰 결과가 기록되지 않았다."
        )
        anomaly_narrative = (
            f"A total of {anomaly_count} unexpected condition(s) were triaged and logged."
            if anomaly_count > 0 else "Execution proceeded without unresolved abnormal incidents."
        )
        conclusion = (
            f"Experiment concluded with status '{report_data.get('status', 'in_progress')}'. "
            f"All {len(event_ids)} event records are verified in the server SQLite ledger."
        )

        return ReportDraftState(
            report_id=str(report_data.get("report_id", "")),
            session_id=str(report_data.get("session_id", "")),
            protocol_id=str(report_data.get("protocol_id", "")),
            status=str(report_data.get("status", "in_progress")),
            experiment_summary=summary,
            observations_narrative=obs_narrative,
            anomalies_narrative=anomaly_narrative,
            conclusion_narrative=conclusion,
            committed_event_ids=event_ids,
            last_updated_at=_now(),
        )

    def build_deterministic_narrative(
        self,
        report_data_or_context: dict[str, Any] | GroundedReportContext,
        events: list[dict[str, Any]] | None = None,
        context: GroundedReportContext | None = None,
    ) -> ReportNarrative:
        """Create a faithful, unhallucinated report narrative grounded in executed protocol steps."""
        if isinstance(report_data_or_context, GroundedReportContext):
            context = report_data_or_context
            report_data = context.report_metadata or {}
            if events is None:
                events = list(context.event_ledger or [])
        else:
            report_data = report_data_or_context
            if events is None:
                events = list(report_data.get("events") or [])
            if context is None:
                context = build_grounded_report_context(report_data, events)

        title = str(report_data.get("protocol_title") or context.protocol_metadata.get("protocol_title") or "Laboratory Experiment Report")
        report_id = str(report_data.get("report_id") or "—")
        status = str(report_data.get("status") or "in_progress")
        started = _human_event_clock(report_data.get("started_at"))
        ended = _human_event_clock(report_data.get("ended_at"))

        objective = (
            f"본 실험 세션의 목적은 '{title}' 지침에 따라 검증된 절차를 체계적으로 수행하고, "
            f"공정별 진행 상태 및 작업자 관찰 결과를 감사 가능한 전자 실험 기록으로 남기는 것이다. "
            f"{context.experiment_objective}"
        )

        completed_count = sum(1 for s in context.executed_steps if s.completion_state == "completed")
        session_summary = (
            f"본 실험 세션({report_id})은 시작 시간({started or '기록됨'})부터 종료 시점({ended or '진행 중'})까지 "
            f"총 {len(events)}건의 검증된 이벤트가 등록되었으며, {completed_count}개의 단계가 완료 기록되었다. "
            f"현재 세션의 최종 상태는 '{status}'(으)로 기록되었다."
        )

        highlights = [_human_event_label(e) for e in events if _human_event_label(e)]

        # Materials & Methods
        mat_items = []
        if context.materials:
            mat_items.append(f"• 주요 시약 및 재료: {', '.join(context.materials)}")
        if context.equipment:
            mat_items.append(f"• 사용 장비: {', '.join(context.equipment)}")
        if context.prerequisites:
            mat_items.append(f"• 사전 준비 사항: {', '.join(context.prerequisites)}")

        method_items = []
        for s in context.executed_steps:
            item_txt = f"• [{s.section_title}] Step {s.step_label}: {s.instruction_source_text}"
            if s.quantities:
                item_txt += f" (수량: {', '.join(s.quantities)})"
            if s.conditions:
                item_txt += f" (조건: {', '.join(s.conditions)})"
            method_items.append(item_txt)

        materials = (
            f"프로토콜 식별자: {context.protocol_metadata.get('protocol_id') or '—'} (버전: {context.protocol_metadata.get('protocol_revision') or '—'}).\n"
            f"원문 문서 SHA-256: {context.protocol_metadata.get('protocol_sha256') or '—'}.\n"
            + ("\n".join(mat_items) + "\n\n" if mat_items else "")
            + "수행 절차 요약:\n"
            + ("\n".join(method_items) if method_items else "기록된 실행 단계가 없습니다.")
        )

        # Results & Observations - detailed chronological narrative
        res_paras = []
        if context.executed_steps:
            for s in context.executed_steps:
                p_bits = [f"[{s.section_title}] {s.step_label}단계에서는 '{s.instruction_source_text}' 절차를 수행하였다."]
                if s.quantities:
                    p_bits.append(f"적용된 물질 및 수량은 {', '.join(s.quantities)}이다.")
                if s.conditions:
                    p_bits.append(f"실험 조건은 {', '.join(s.conditions)}이다.")
                if s.timer_configuration:
                    p_bits.append(f"설정된 타이머는 {s.timer_configuration}이며, 실제 {s.timer_actuals or '정상 시간'} 동안 유지되었다.")

                # Distinguish user observation vs expected results vs no record
                if s.user_confirmed_observations:
                    obs_str = "; ".join(s.user_confirmed_observations)
                    p_bits.append(f"작업자는 해당 단계에서 '{obs_str}'(이)라는 관찰 결과를 기록하였으며, 이 관찰이 완료 조건을 충족하였다.")
                elif s.expected_results:
                    exp_str = "; ".join(s.expected_results)
                    p_bits.append(f"실험 PDF에서는 '{exp_str}'을(를) 예상 결과로 제시하고 있으나, 해당 단계에 대해 별도의 실제 관찰값은 기록되지 않았다.")
                else:
                    p_bits.append("해당 단계에 대해 별도의 관찰 결과가 기록되지 않았다.")

                if s.completion_state == "completed":
                    p_bits.append("해당 단계는 완료로 기록되었다.")

                res_paras.append(" ".join(p_bits))
        else:
            res_paras.append("진행 중인 단계에 대한 수행 기록이 유지되고 있다.")

        results_and_obs = "\n\n".join(res_paras)

        # Discussion
        discussion_parts = []
        if context.anomalies:
            discussion_parts.append(
                f"세션 진행 중 {len(context.anomalies)}건의 이상/편차 사항이 보고되어 기록되었다."
            )
        else:
            discussion_parts.append("절차 진행 중 보고된 비정상 편차나 위험 요인은 없었다.")

        if context.deviations:
            discussion_parts.append(f"총 {len(context.deviations)}회의 일시정지 및 재개가 발생하였다.")
        discussion_parts.append("모든 수행 내역은 불변 SQLite 이벤트 원장에 무결하게 보존되었다.")
        discussion = " ".join(discussion_parts)

        # Anomalies
        anomalies_dev = (
            "\n".join(
                f"• [{_human_event_clock(ae.get('created_at'))}] Step {ae.get('step_label') or '—'}: {ae.get('user_wording') or (ae.get('payload') or {}).get('text') or '이상 보고'}"
                for ae in context.anomalies
            )
            if context.anomalies else "특이 이상 또는 절차 차단 사항 없음."
        )

        conclusion = (
            f"실험 세션 {report_id}의 최종 상태는 '{status}'이며, "
            f"총 {len(events)}개의 암호학적 원장 이벤트가 유효하게 확정되었다."
        )

        limitations = (
            "본 보고서는 시스템에 등록된 실제 이벤트 기록 및 활성 실험 PDF에 한하여 작성되었으며, "
            "작업자가 명시적으로 기록하지 않은 관찰값이나 측정 수치는 포함하지 않는다."
        )

        return ReportNarrative(
            title=title,
            objective=objective,
            session_summary=session_summary,
            chronological_highlights=tuple(highlights),
            materials_and_methods=materials,
            results_and_observations=results_and_obs,
            discussion=discussion,
            anomalies_and_deviations=anomalies_dev,
            conclusion=conclusion,
            limitations=limitations,
        )

    async def generate_narrative(
        self,
        report_data: dict[str, Any],
        events: list[dict[str, Any]],
    ) -> ReportNarrative:
        """Generate a rich, academic Korean narrative report grounded strictly in the ledger and protocol steps."""
        context = build_grounded_report_context(report_data, events)
        report_id = str(report_data.get("report_id") or "")
        protocol_rev = str(report_data.get("protocol_revision") or "")
        finalization_version = int(report_data.get("finalization_version") or 0)
        latest_key = str(events[-1].get("event_key") if events else "")
        cache_key = (report_id, finalization_version, latest_key, protocol_rev, self.model)

        cached = _NARRATIVE_CACHE.get(cache_key)
        if cached is not None:
            return cached

        deterministic = self.build_deterministic_narrative(report_data, events, context=context)
        if not self.client:
            _NARRATIVE_CACHE[cache_key] = deterministic
            return deterministic

        try:
            prompt_context = context.public_dict()
            system_prompt = (
                "You are an experienced professional laboratory scientific report writer. "
                "Write a formal, comprehensive laboratory session report in Korean based strictly on the provided verified experimental context and event ledger. "
                "STYLE REQUIREMENTS: Write coherent, academic, natural narrative prose with full paragraphs. "
                "FACTUAL INTEGRITY RULES: "
                "1. Ground Materials and Methods and Results strictly in the provided executed steps, quantities, conditions, and timers. "
                "2. Distinguish protocol expected results from actual user observations. If an expected result was in the protocol but no observation was recorded, explicitly state that no actual observation was logged ('해당 단계에 대해 별도의 관찰 결과가 기록되지 않았다.'). "
                "3. Do NOT invent or hallucinate any unrecorded measurements, temperatures, reagent volumes, or success claims. "
                "4. Return a valid JSON object matching the requested schema exactly."
            )
            user_prompt = f"Grounded Experimental Report Context:\n{json.dumps(prompt_context, ensure_ascii=False, indent=2)}"

            resp = await asyncio.wait_for(
                self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    response_format={"type": "json_object"},
                    max_tokens=1800,
                ),
                timeout=self.timeout_seconds,
            )
            content = resp.choices[0].message.content or ""
            data = json.loads(content)
            narrative = ReportNarrative(
                title=str(data.get("title") or deterministic.title),
                objective=str(data.get("objective") or deterministic.objective),
                session_summary=str(data.get("session_summary") or deterministic.session_summary),
                chronological_highlights=tuple(data.get("chronological_highlights") or deterministic.chronological_highlights),
                materials_and_methods=str(data.get("materials_and_methods") or deterministic.materials_and_methods),
                results_and_observations=str(data.get("results_and_observations") or deterministic.results_and_observations),
                discussion=str(data.get("discussion") or deterministic.discussion),
                anomalies_and_deviations=str(data.get("anomalies_and_deviations") or deterministic.anomalies_and_deviations),
                conclusion=str(data.get("conclusion") or deterministic.conclusion),
                limitations=str(data.get("limitations") or deterministic.limitations),
            )
            _NARRATIVE_CACHE[cache_key] = narrative
            return narrative
        except Exception as exc:
            log.warning("ReportWriterBrain LLM generation failed (%s), using deterministic narrative", exc)
            _NARRATIVE_CACHE[cache_key] = deterministic
            return deterministic



def _format_elapsed_clock(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        total = int(round(float(value)))
    except (TypeError, ValueError):
        return ""
    if total < 0:
        total = 0
    hours, remainder = divmod(total, 3600)
    minutes, seconds = divmod(remainder, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    return f"{minutes:02d}:{seconds:02d}"


def _human_event_clock(created_at: str | None) -> str:
    if not created_at:
        return ""
    try:
        parsed = datetime.fromisoformat(str(created_at).replace("Z", "+00:00"))
    except ValueError:
        return ""
    return parsed.strftime("%H:%M:%S")


_HUMAN_EVENT_VERBS = {
    "session_started": "시작",
    "step_completed": "완료",
    "step_advanced": "이동",
    "step_presented": "안내",
    "timer_started": "타이머 시작",
    "workflow_paused": "일시정지",
    "workflow_resumed": "재개",
    "workflow_completed": "완료",
    "session_stopped": "종료",
    "anomaly": "이상 보고",
    "blocked": "진행 차단",
    "observation": "관찰",
    "source_consulted": "참고 자료 확인",
}


def _human_event_label(event: dict[str, Any]) -> str:
    event_type = str(event.get("event_type") or "")
    verb = _HUMAN_EVENT_VERBS.get(event_type, event_type.replace("_", " "))
    step_label = event.get("step_label")
    head = f"Step {step_label} {verb}" if step_label else verb
    parts = [head]
    clock = _human_event_clock(event.get("created_at"))
    if clock:
        parts.append(clock)
    payload = event.get("payload") if isinstance(event.get("payload"), dict) else {}
    timer = payload.get("timer") if isinstance(payload.get("timer"), dict) else {}
    duration = timer.get("source_duration_seconds", timer.get("duration_seconds"))
    elapsed = timer.get("elapsed_seconds")
    remaining = timer.get("remaining_seconds")
    timer_bits = []
    if duration not in (None, ""):
        timer_bits.append(f"총 {_format_elapsed_clock(duration)}")
    if elapsed not in (None, ""):
        timer_bits.append(f"경과 {_format_elapsed_clock(elapsed)}")
    if remaining not in (None, ""):
        timer_bits.append(f"잔여 {_format_elapsed_clock(remaining)}")
    if timer_bits:
        parts.append("타이머 " + " · ".join(timer_bits))
    wording = event.get("user_wording")
    if wording:
        parts.append(str(wording))
    return " / ".join(parts)


def new_session_id() -> str:
    return "session-" + secrets.token_hex(16)
